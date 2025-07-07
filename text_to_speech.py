"""
Formatmaskinen - Text-to-Speech och Podcast-generator med LiteLLM
Använder LiteLLM proxy istället för direkt OpenAI API
"""

# === IMPORTS ===
from flask import Flask, request, jsonify, send_file, Response
from flask_cors import CORS
from pydantic import BaseModel, ValidationError
from typing import Optional, Dict, Any, List
import io
import os
import tempfile
import shutil
import re
from pathlib import Path
from dotenv import load_dotenv
import asyncio
import httpx
import json
import litellm

# PDF/DOCX imports
import PyPDF2
import pdfplumber
from docx import Document
import chardet
import mimetypes

# Web scraping imports
import requests
from bs4 import BeautifulSoup
from urllib.parse import urlparse, urljoin

# === CONFIGURATION ===

# Ladda miljövariabler
env_path = Path(__file__).parent / '.env'
load_dotenv(str(env_path))

# LiteLLM configuration
LITELLM_API_KEY = os.getenv('LITELLM_API_KEY')
LITELLM_BASE_URL = os.getenv('LITELLM_BASE_URL', 'https://anast.ita.chalmers.se:4000')

# Configure LiteLLM
litellm.api_base = LITELLM_BASE_URL
litellm.drop_params = True

# Flask app
app = Flask(__name__)
CORS(app)

# Temp-mapp för uppladdade filer
UPLOAD_DIR = Path(__file__).parent / "uploads"
UPLOAD_DIR.mkdir(exist_ok=True)


# === FILE HANDLER CLASS ===

class FileHandler:
    """Hanterar olika filformat och textextraktion"""
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self.extract_from_pdf,
            '.docx': self.extract_from_docx,
            '.doc': self.extract_from_doc,
            '.txt': self.extract_from_txt,
            '.md': self.extract_from_txt,
            '.rtf': self.extract_from_rtf,
        }
    
    def extract_text(self, file_path: str) -> Dict[str, Any]:
        """Huvudmetod för att extrahera text från valfri fil"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"Filen finns inte: {file_path}")
        
        file_ext = file_path.suffix.lower()
        
        if file_ext not in self.supported_formats:
            raise ValueError(f"Filformat '{file_ext}' stöds ej. Stödda format: {', '.join(self.supported_formats.keys())}")
        
        extractor = self.supported_formats[file_ext]
        result = extractor(str(file_path))
        
        result['file_name'] = file_path.name
        result['file_type'] = file_ext
        result['file_size_bytes'] = file_path.stat().st_size
        result['file_size_mb'] = round(result['file_size_bytes'] / (1024 * 1024), 2)
        
        return result
    
    def extract_from_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extrahera text från PDF"""
        text = ""
        metadata = {}
        
        try:
            with pdfplumber.open(file_path) as pdf:
                metadata['pages'] = len(pdf.pages)
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
        except:
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    metadata['pages'] = len(pdf_reader.pages)
                    
                    for page_num in range(metadata['pages']):
                        page = pdf_reader.pages[page_num]
                        text += page.extract_text() + "\n\n"
            except Exception as e:
                raise Exception(f"Kunde inte läsa PDF: {str(e)}")
        
        text = self.clean_text(text)
        paragraphs = self.split_into_paragraphs(text)
        
        return {
            "text": text,
            "paragraphs": paragraphs,
            "word_count": len(text.split()),
            "character_count": len(text),
            "metadata": metadata
        }
    
    def extract_from_docx(self, file_path: str) -> Dict[str, Any]:
        """Extrahera text från DOCX"""
        try:
            doc = Document(file_path)
            paragraphs = []
            full_text = ""
            
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())
                    full_text += para.text + "\n\n"
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text += cell.text + "\n"
                full_text += "\n"
            
            text = self.clean_text(full_text)
            
            metadata = {
                "paragraphs_count": len(paragraphs),
                "tables_count": len(doc.tables),
            }
            
            try:
                core_props = doc.core_properties
                metadata.update({
                    "author": core_props.author or "Unknown",
                    "title": core_props.title or "Untitled",
                    "created": str(core_props.created) if core_props.created else None,
                    "modified": str(core_props.modified) if core_props.modified else None,
                })
            except:
                pass
            
            return {
                "text": text,
                "paragraphs": self.split_into_paragraphs(text),
                "word_count": len(text.split()),
                "character_count": len(text),
                "metadata": metadata
            }
            
        except Exception as e:
            raise Exception(f"Kunde inte läsa DOCX: {str(e)}")
    
    def extract_from_doc(self, file_path: str) -> Dict[str, Any]:
        """Extrahera text från DOC (legacy Word)"""
        raise NotImplementedError(
            "DOC-format (legacy Word) stöds inte ännu. "
            "Spara filen som DOCX i Word först."
        )
    
    def extract_from_txt(self, file_path: str) -> Dict[str, Any]:
        """Extrahera text från TXT/MD-filer"""
        try:
            with open(file_path, 'rb') as file:
                raw_data = file.read()
                detected = chardet.detect(raw_data)
                encoding = detected['encoding'] or 'utf-8'
            
            with open(file_path, 'r', encoding=encoding) as file:
                text = file.read()
            
            if file_path.endswith('.md'):
                text = re.sub(r'#{1,6}\s*', '', text)
                text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
                text = re.sub(r'\*([^*]+)\*', r'\1', text)
                text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
                text = re.sub(r'`([^`]+)`', r'\1', text)
                text = re.sub(r'```[^`]*```', '', text)
            
            text = self.clean_text(text)
            
            return {
                "text": text,
                "paragraphs": self.split_into_paragraphs(text),
                "word_count": len(text.split()),
                "character_count": len(text),
                "metadata": {
                    "encoding": encoding,
                    "confidence": detected.get('confidence', 0)
                }
            }
            
        except Exception as e:
            raise Exception(f"Kunde inte läsa textfil: {str(e)}")
    
    def extract_from_rtf(self, file_path: str) -> Dict[str, Any]:
        """Extrahera text från RTF"""
        raise NotImplementedError(
            "RTF-format stöds inte ännu. "
            "Konvertera till DOCX eller TXT först."
        )
    
    def clean_text(self, text: str) -> str:
        """Rensa upp extraherad text"""
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'[^\w\s\-–—.,;:!?()"\'åäöÅÄÖéÉèÈüÜ\n]', '', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = text.strip()
        return text
    
    def split_into_paragraphs(self, text: str, max_length: int = 1000) -> List[str]:
        """Dela upp text i lagom stora stycken för TTS"""
        paragraphs = text.split('\n\n')
        result = []
        current_chunk = ""
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
                
            if len(para) > max_length:
                sentences = re.split(r'(?<=[.!?])\s+', para)
                
                for sentence in sentences:
                    if len(current_chunk) + len(sentence) + 1 < max_length:
                        current_chunk += sentence + " "
                    else:
                        if current_chunk:
                            result.append(current_chunk.strip())
                        current_chunk = sentence + " "
            else:
                if len(current_chunk) + len(para) + 2 < max_length:
                    if current_chunk:
                        current_chunk += "\n\n"
                    current_chunk += para
                else:
                    if current_chunk:
                        result.append(current_chunk.strip())
                    current_chunk = para
        
        if current_chunk:
            result.append(current_chunk.strip())
        
        return result


# === URL HANDLER CLASS ===

class URLHandler:
    """Hanterar URL:er och extraherar text från webbsidor och PDF:er"""
    
    def __init__(self):
        self.file_handler = FileHandler()
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        })
        
        self.content_selectors = [
            'article', '[role="main"]', '.content', '.post-content',
            '.entry-content', '.article-content', '.main-content',
            'main', '.text', '.body'
        ]
    
    def extract_from_url(self, url: str, timeout: int = 30) -> Dict[str, Any]:
        """Huvudmetod för att extrahera text från valfri URL"""
        parsed_url = urlparse(url)
        if not parsed_url.scheme or not parsed_url.netloc:
            raise ValueError(f"Ogiltig URL: {url}")
        
        try:
            head_response = self.session.head(url, timeout=timeout, allow_redirects=True)
            content_type = head_response.headers.get('content-type', '').lower()
            
            if 'application/pdf' in content_type or url.lower().endswith('.pdf'):
                return self._extract_pdf_from_url(url, timeout)
            
            return self._extract_webpage_content(url, timeout)
            
        except requests.RequestException as e:
            raise Exception(f"Kunde inte hämta URL {url}: {str(e)}")
    
    def _extract_pdf_from_url(self, url: str, timeout: int) -> Dict[str, Any]:
        """Ladda ner PDF från URL och extrahera text"""
        try:
            response = self.session.get(url, timeout=timeout, stream=True)
            response.raise_for_status()
            
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
                temp_path = temp_file.name
                
                for chunk in response.iter_content(chunk_size=8192):
                    if chunk:
                        temp_file.write(chunk)
            
            try:
                result = self.file_handler.extract_text(temp_path)
                result['source_url'] = url
                result['source_type'] = 'pdf_url'
                result['content_type'] = response.headers.get('content-type')
                return result
                
            finally:
                try:
                    os.unlink(temp_path)
                except:
                    pass
                    
        except Exception as e:
            raise Exception(f"Kunde inte hämta PDF från {url}: {str(e)}")
    
    def _extract_webpage_content(self, url: str, timeout: int) -> Dict[str, Any]:
        """Extrahera text från webbsida"""
        try:
            response = self.session.get(url, timeout=timeout)
            response.raise_for_status()
            
            if response.encoding is None:
                response.encoding = response.apparent_encoding
            
            soup = BeautifulSoup(response.text, 'lxml')
            
            for script in soup(["script", "style", "nav", "header", "footer", "aside"]):
                script.decompose()
            
            main_content = self._find_main_content(soup)
            
            if not main_content:
                main_content = soup.find('body') or soup
            
            text = self._extract_clean_text(main_content)
            
            title_tag = soup.find('title')
            title = title_tag.get_text().strip() if title_tag else urlparse(url).netloc
            
            meta_desc = soup.find('meta', attrs={'name': 'description'})
            description = meta_desc.get('content', '').strip() if meta_desc else ''
            
            text = self._clean_webpage_text(text)
            paragraphs = self._split_into_paragraphs(text)
            
            return {
                "text": text,
                "paragraphs": paragraphs,
                "word_count": len(text.split()),
                "character_count": len(text),
                "source_url": url,
                "source_type": "webpage",
                "metadata": {
                    "title": title,
                    "description": description,
                    "content_type": response.headers.get('content-type'),
                    "final_url": response.url,
                    "status_code": response.status_code
                }
            }
            
        except Exception as e:
            raise Exception(f"Kunde inte extrahera text från {url}: {str(e)}")
    
    def _find_main_content(self, soup: BeautifulSoup) -> Optional[Any]:
        """Försök hitta huvudinnehållet på webbsidan"""
        for selector in self.content_selectors:
            content = soup.select_one(selector)
            if content and len(content.get_text().strip()) > 100:
                return content
        
        text_containers = soup.find_all(['div', 'section', 'article'])
        if text_containers:
            text_containers.sort(key=lambda x: len(x.get_text()), reverse=True)
            
            for container in text_containers[:3]:
                text_length = len(container.get_text().strip())
                if text_length > 200:
                    return container
        
        return None
    
    def _extract_clean_text(self, element) -> str:
        """Extrahera ren text från HTML-element"""
        for tag in element.find_all(['p', 'div', 'br', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li']):
            tag.insert_after('\n')
        
        text = element.get_text()
        return text
    
    def _clean_webpage_text(self, text: str) -> str:
        """Rensa upp extraherad webbtext"""
        text = re.sub(r'\n+', '\n', text)
        text = re.sub(r'[ \t]+', ' ', text)
        
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        
        content_lines = []
        for line in lines:
            if len(line) > 20 or '.' in line:
                content_lines.append(line)
        
        text = '\n\n'.join(content_lines)
        text = re.sub(r'[^\w\s\-–—.,;:!?()"\'åäöÅÄÖéÉèÈüÜ\n]', '', text)
        text = text.strip()
        
        return text
    
    def _split_into_paragraphs(self, text: str, max_length: int = 1000) -> list:
        """Dela upp text i lagom stora stycken"""
        paragraphs = text.split('\n\n')
        result = []
        current_chunk = ""
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
                
            if len(para) > max_length:
                sentences = re.split(r'(?<=[.!?])\s+', para)
                
                for sentence in sentences:
                    if len(current_chunk) + len(sentence) + 1 < max_length:
                        current_chunk += sentence + " "
                    else:
                        if current_chunk:
                            result.append(current_chunk.strip())
                        current_chunk = sentence + " "
            else:
                if len(current_chunk) + len(para) + 2 < max_length:
                    if current_chunk:
                        current_chunk += "\n\n"
                    current_chunk += para
                else:
                    if current_chunk:
                        result.append(current_chunk.strip())
                    current_chunk = para
        
        if current_chunk:
            result.append(current_chunk.strip())
        
        return result
    
    def validate_url(self, url: str) -> Dict[str, Any]:
        """Validera URL utan att hämta innehåll"""
        try:
            parsed = urlparse(url)
            
            if not parsed.scheme:
                url = 'https://' + url
                parsed = urlparse(url)
            
            if not parsed.netloc:
                raise ValueError("Ogiltig URL struktur")
            
            try:
                response = self.session.head(url, timeout=10, allow_redirects=True)
                content_type = response.headers.get('content-type', '').lower()
                
                if 'application/pdf' in content_type or url.lower().endswith('.pdf'):
                    url_type = 'pdf'
                elif 'text/html' in content_type:
                    url_type = 'webpage'
                else:
                    url_type = 'unknown'
                
                return {
                    "valid": True,
                    "url": response.url,
                    "status_code": response.status_code,
                    "content_type": content_type,
                    "url_type": url_type,
                    "accessible": response.status_code == 200
                }
                
            except requests.RequestException:
                return {
                    "valid": True,
                    "url": url,
                    "accessible": False,
                    "error": "Kunde inte nå URL:en"
                }
                
        except Exception as e:
            return {
                "valid": False,
                "url": url,
                "error": str(e)
            }


# === TEXT TO SPEECH CLASS ===

class TextToSpeech:
    """Text-to-Speech handler med ElevenLabs"""
    
    VOICE_IDS = {
        "charlotte": "XB0fDUnXU5powFXDhCwa",
        "sarah": "EXAVITQu4vr4xnSDxMaL",
        "george": "JBFqnCBsd6RMkjVDRZzb",
        "charlie": "IKne3meq5aSn9XLyUdCD",
        "aria": "9BWtsMINqrJLrRacOk9x",
        "laura": "FGY2WhTYpPnrIDTdsKH5",
        "river": "SAz9YHcvj6GT2YYXdXww",
        "liam": "TX3LPaxmHKxFdv7VOQHJ",
        "alice": "Xb7hH8MSUJpSbSDYk0k2",
        "callum": "N2lVS1w4EtoT3dr4eOWO",
        "default": "EXAVITQu4vr4xnSDxMaL",
        "astrid": "XB0fDUnXU5powFXDhCwa",
    }
    
    CONVERSATION_PAIRS = {
        "host_guest_en": {
            "host": "21m00Tcm4TlvDq8ikWAM",
            "guest": "JBFqnCBsd6RMkjVDRZzb"
        },
        "interview_sv": {
            "interviewer": "IKne3meq5aSn9XLyUdCD",
            "interviewee": "21m00Tcm4TlvDq8ikWAM"
        },
        "debate": {
            "speaker1": "XB0fDUnXU5powFXDhCwa",
            "speaker2": "JBFqnCBsd6RMkjVDRZzb"
        }
    }
    
    SUPPORTED_LANGUAGES = [
        "sv", "en", "de", "fr", "es", "it", "pt", "pl", 
        "tr", "ru", "nl", "cs", "ar", "zh", "ja", "ko"
    ]
    
    AVAILABLE_MODELS = {
        "eleven_multilingual_v2": "Standard multilingual model",
        "eleven_turbo_v2": "Fast, low-latency model",
        "eleven_multilingual_v1": "Legacy multilingual model",
        "eleven_monolingual_v1": "English-only model", 
        "eleven_turbo_v2_5": "Ultra-fast model (75ms latency)",
        "eleven_v3_alpha": "Most realistic model (alpha - may require access)",
        "eleven_v3": "Most realistic model (if available)"
    }
    
    def __init__(self, api_key: Optional[str] = None):
        """Initiera TTS-modul"""
        self.api_key = api_key or os.getenv("ELEVENLABS_API_KEY")
        if not self.api_key:
            raise ValueError("ELEVENLABS_API_KEY saknas!")
        
        # LiteLLM används istället för OpenAI direkt
        self.litellm_api_key = LITELLM_API_KEY
        self.litellm_available = bool(self.litellm_api_key)
        
        if not self.litellm_available:
            print("⚠️ LITELLM_API_KEY saknas - dialogue generation kommer att vara begränsad")
        
        self._voice_cache = {}
        self._conversation_templates = self._create_conversation_templates()
    
    def convert_text_to_speech(
        self,
        text: str,
        voice: str = "astrid",
        language: str = "sv",
        model: str = "eleven_multilingual_v2"
    ) -> bytes:
        """Konvertera text till tal (synkron wrapper för async-funktionen)"""
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        try:
            return loop.run_until_complete(
                self.convert_text_to_speech_async(text, voice, language, model)
            )
        finally:
            loop.close()
    
    async def convert_text_to_speech_async(
        self,
        text: str,
        voice: str = "astrid",
        language: str = "sv",
        model: str = "eleven_multilingual_v2"
    ) -> bytes:
        """Konvertera text till tal (asynkron)"""
        if not text or not text.strip():
            raise ValueError("Text kan inte vara tom")
        if language not in self.SUPPORTED_LANGUAGES:
            raise ValueError(f"Språk '{language}' stöds ej. Använd: {', '.join(self.SUPPORTED_LANGUAGES)}")
        
        voice_id = self.VOICE_IDS.get(voice, self.VOICE_IDS["default"])
        api_key = self.api_key
        headers = {
            "xi-api-key": api_key,
            "Content-Type": "application/json"
        }
        url = f"https://api.elevenlabs.io/v1/text-to-speech/{voice_id}"
        data = {
            "text": text,
            "model_id": model,
            "voice_settings": {},
            "output_format": "mp3_44100_128"
        }
        
        async with httpx.AsyncClient(timeout=60) as client:
            response = await client.post(url, headers=headers, json=data)
            if response.status_code != 200:
                raise Exception(f"TTS API error: {response.status_code} {response.text}")
            return response.content
    
    def list_available_voices(self) -> Dict[str, Any]:
        """Lista alla tillgängliga röster från ElevenLabs API (synkron wrapper)"""
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        try:
            return loop.run_until_complete(self.list_available_voices_async())
        finally:
            loop.close()
    
    async def list_available_voices_async(self) -> Dict[str, Any]:
        """Lista alla tillgängliga röster från ElevenLabs API (asynkron)"""
        if not self._voice_cache:
            try:
                api_key = self.api_key
                headers = {"xi-api-key": api_key}
                url = "https://api.elevenlabs.io/v1/voices"
                
                async with httpx.AsyncClient(timeout=30) as client:
                    response = await client.get(url, headers=headers)
                    if response.status_code == 200:
                        voices_data = response.json()
                        self._voice_cache = {}
                        
                        for voice in voices_data.get("voices", []):
                            voice_id = voice.get("voice_id")
                            voice_name = voice.get("name")
                            labels = voice.get("labels", {})
                            
                            gender = labels.get("gender", "unknown")
                            accent = labels.get("accent", "")
                            age = labels.get("age", "")
                            descriptive = labels.get("descriptive", "")
                            
                            description_parts = []
                            if gender != "unknown":
                                description_parts.append(gender)
                            if accent:
                                description_parts.append(accent)
                            if age:
                                description_parts.append(age)
                            if descriptive:
                                description_parts.append(descriptive)
                            
                            description = ", ".join(description_parts) if description_parts else "General voice"
                            
                            self._voice_cache[voice_name] = {
                                "id": voice_id,
                                "description": description,
                                "labels": labels,
                                "preview_url": voice.get("preview_url", "")
                            }
                    else:
                        print(f"API fel vid hämtning av röster: {response.status_code}")
                        self._voice_cache = self._create_fallback_voices()
                        
            except Exception as e:
                print(f"Kunde inte hämta röster från API: {e}")
                self._voice_cache = self._create_fallback_voices()
        
        return {
            "voices": self._voice_cache,
            "languages": self.SUPPORTED_LANGUAGES,
            "default_voice": "charlotte"
        }
    
    def _create_fallback_voices(self) -> Dict[str, Any]:
        """Skapa fallback röster med korrekta beskrivningar"""
        return {
            "Charlotte": {
                "id": "XB0fDUnXU5powFXDhCwa",
                "description": "female, swedish, young, relaxed",
                "labels": {"gender": "female", "accent": "swedish", "age": "young"}
            },
            "Sarah": {
                "id": "EXAVITQu4vr4xnSDxMaL", 
                "description": "female, american, young, professional",
                "labels": {"gender": "female", "accent": "american", "age": "young"}
            },
            "George": {
                "id": "JBFqnCBsd6RMkjVDRZzb",
                "description": "male, british, middle_aged, mature", 
                "labels": {"gender": "male", "accent": "british", "age": "middle_aged"}
            },
            "Charlie": {
                "id": "IKne3meq5aSn9XLyUdCD",
                "description": "male, australian, young, hyped",
                "labels": {"gender": "male", "accent": "australian", "age": "young"}
            },
            "Aria": {
                "id": "9BWtsMINqrJLrRacOk9x",
                "description": "female, american, middle_aged, husky",
                "labels": {"gender": "female", "accent": "american", "age": "middle_aged"}
            },
            "River": {
                "id": "SAz9YHcvj6GT2YYXdXww",
                "description": "neutral, american, middle_aged, calm",
                "labels": {"gender": "neutral", "accent": "american", "age": "middle_aged"}
            }
        }
    
    def _create_conversation_templates(self) -> Dict[str, Dict[str, Any]]:
        """Skapa conversation templates för olika podcast-stilar"""
        return {
            "tech_interview": {
                "name": "Tech Interview",
                "description": "Deep dive technical discussion",
                "host_persona": "Experienced tech journalist who asks insightful questions",
                "guest_persona": "Subject matter expert who explains complex topics clearly",
                "style": "professional, curious, educational",
                "opening": "Welcome to our tech deep dive! Today we're exploring an fascinating topic.",
                "closing": "Thank you for this enlightening discussion. That's all for today!"
            },
            "casual_chat": {
                "name": "Casual Chat",
                "description": "Friendly conversation between peers", 
                "host_persona": "Friendly host who loves learning new things",
                "guest_persona": "Knowledgeable friend sharing insights",
                "style": "casual, friendly, conversational",
                "opening": "Hey everyone! Today I'm chatting with someone about a really interesting topic.",
                "closing": "This was such a great conversation! Thanks for joining us today."
            },
            "news_analysis": {
                "name": "News Analysis",
                "description": "Analytical discussion of current events",
                "host_persona": "Serious news anchor who provides context",
                "guest_persona": "Expert analyst offering deep insights",
                "style": "professional, analytical, informative",
                "opening": "Welcome to our analysis segment. Let's examine today's important developments.",
                "closing": "That concludes our analysis. We'll continue following this story."
            },
            "educational": {
                "name": "Educational",
                "description": "Teaching-focused content explanation",
                "host_persona": "Curious student asking great questions",
                "guest_persona": "Patient teacher explaining concepts step by step",
                "style": "educational, clear, methodical",
                "opening": "Today we're going to learn about an important topic. Let's dive in!",
                "closing": "I hope this helped you understand the topic better. Keep learning!"
            },
            "storytelling": {
                "name": "Storytelling",
                "description": "Narrative-driven content presentation",
                "host_persona": "Engaging storyteller who sets the scene",
                "guest_persona": "Co-narrator who adds details and perspectives",
                "style": "narrative, engaging, dramatic",
                "opening": "Let me tell you a story that will change how you think about this topic.",
                "closing": "And that's how our story ends. What did you think of that journey?"
            }
        }
    
    def estimate_cost(self, text: str) -> Dict[str, float]:
        """Uppskatta kostnad för TTS-konvertering"""
        char_count = len(text)
        estimated_cost = char_count * 0.00018
        
        return {
            "characters": char_count,
            "estimated_cost_usd": round(estimated_cost, 4),
            "estimated_cost_sek": round(estimated_cost * 10.5, 2)
        }
    
    def generate_enhanced_conversation(
        self,
        text: str,
        template_style: str = "casual_chat",
        host_voice_id: str = "EXAVITQu4vr4xnSDxMaL",
        guest_voice_id: str = "JBFqnCBsd6RMkjVDRZzb",
        language: str = "en",
        custom_prompt: str = "",
        model: str = "eleven_multilingual_v2"
    ) -> bytes:
        """Generera förbättrad conversation med LiteLLM + ElevenLabs (synkron wrapper)"""
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        try:
            return loop.run_until_complete(
                self.generate_enhanced_conversation_async(
                    text, template_style, host_voice_id, guest_voice_id, 
                    language, custom_prompt, model
                )
            )
        finally:
            loop.close()
    
    async def generate_enhanced_conversation_async(
        self,
        text: str,
        template_style: str = "casual_chat",
        host_voice_id: str = "EXAVITQu4vr4xnSDxMaL",
        guest_voice_id: str = "JBFqnCBsd6RMkjVDRZzb",
        language: str = "en",
        custom_prompt: str = "",
        model: str = "eleven_multilingual_v2"
    ) -> bytes:
        """Generera förbättrad conversation med LiteLLM + ElevenLabs (asynkron)"""
        print(f"🎙️ Genererar förbättrad conversation med stil: {template_style}")
        
        template = self._conversation_templates.get(template_style, self._conversation_templates["casual_chat"])
        dialogue_segments = self._generate_ai_dialogue(text, template, language, custom_prompt)
        
        audio_segments = []
        for i, segment in enumerate(dialogue_segments):
            speaker = segment["speaker"]
            speech_text = segment["text"]
            
            if speaker == "host":
                voice_id = host_voice_id
            else:
                voice_id = guest_voice_id
            
            print(f"🎯 Genererar audio för {speaker}: {speech_text[:50]}...")
            
            audio_data = await self.convert_text_to_speech_with_voice_id(
                text=speech_text,
                voice_id=voice_id,
                language=language,
                model=model
            )
            
            audio_segments.append(audio_data)
            
            if i < len(dialogue_segments) - 1:
                pause_audio = await self._generate_pause(0.5)
                audio_segments.append(pause_audio)
        
        combined_audio = self._combine_audio_segments(audio_segments)
        
        print(f"✅ Conversation genererad! {len(combined_audio)} bytes")
        return combined_audio

    def generate_conversation_from_text(
        self,
        text: str,
        conversation_style: str = "host_guest_en",
        custom_voices: Optional[Dict[str, str]] = None,
        model: str = "eleven_multilingual_v2"
    ) -> bytes:
        """Generera conversational podcast från text med ElevenLabs (synkron wrapper)"""
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        try:
            return loop.run_until_complete(
                self.generate_conversation_from_text_async(
                    text, conversation_style, custom_voices, model
                )
            )
        finally:
            loop.close()
    
    async def generate_conversation_from_text_async(
        self,
        text: str,
        conversation_style: str = "host_guest_en",
        custom_voices: Optional[Dict[str, str]] = None,
        model: str = "eleven_multilingual_v2"
    ) -> bytes:
        """Generera conversational podcast från text med ElevenLabs (asynkron)"""
        if custom_voices:
            voices = custom_voices
        elif conversation_style in self.CONVERSATION_PAIRS:
            voices = self.CONVERSATION_PAIRS[conversation_style]
        else:
            voices = self.CONVERSATION_PAIRS["host_guest_en"]
        
        dialogue_segments = self._structure_text_to_dialogue(text, list(voices.keys()))
        
        audio_segments = []
        for segment in dialogue_segments:
            speaker = segment["speaker"]
            speech_text = segment["text"]
            voice_id = voices.get(speaker, list(voices.values())[0])
            
            audio_data = await self.convert_text_to_speech_with_voice_id(
                text=speech_text,
                voice_id=voice_id,
                language=language if 'language' in locals() else "en",
                model=model
            )
            
            audio_segments.append(audio_data)
        
        combined_audio = self._combine_audio_segments(audio_segments)
        
        return combined_audio
    
    def _generate_ai_dialogue(
        self,
        text: str,
        template: Dict[str, Any],
        language: str,
        custom_prompt: str = ""
    ) -> List[Dict[str, str]]:
        """Generera intelligent dialogue med LiteLLM"""
        if not self.litellm_available:
            return self._structure_text_to_dialogue(text, ["host", "guest"])
        
        system_prompt = self._create_dialogue_prompt(template, language, custom_prompt)
        
        # Models to try on LiteLLM proxy
        models_to_try = [
            "claude-haiku-3.5",
            "claude-sonnet-3.7",
            "gpt-4.1-2025-04-14",
            "gpt-4.5-preview",
            "claude-sonnet-4",
            "o1"
        ]
        
        response = None
        for model in models_to_try:
            try:
                print(f"🤖 Trying model {model} for dialogue generation...")
                
                response = litellm.completion(
                    model=model,
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": f"Please create a natural conversation based on this content:\n\n{text}"}
                    ],
                    api_key=self.litellm_api_key,
                    max_tokens=2000,
                    temperature=0.7,
                    base_url=LITELLM_BASE_URL
                )
                
                print(f"SUCCESS: Model {model} worked for dialogue generation!")
                dialogue_text = response.choices[0].message.content
                return self._parse_ai_dialogue(dialogue_text)
                
            except Exception as e:
                print(f"Model {model} failed: {str(e)}")
                continue
        
        # Fallback if no model works
        print("⚠️ All models failed, using fallback dialogue structure")
        return self._structure_text_to_dialogue(text, ["host", "guest"])
    
    def _create_dialogue_prompt(self, template: Dict[str, Any], language: str, custom_prompt: str) -> str:
        """Skapa LiteLLM prompt för dialogue generation"""
        lang_instruction = "in English" if language == "en" else "på svenska" if language == "sv" else f"in {language}"
        
        base_prompt = f"""You are creating a natural podcast conversation {lang_instruction} between two people:

HOST: {template['host_persona']}
GUEST: {template['guest_persona']}

Style: {template['style']}

Create a natural, engaging dialogue that:
1. Starts with: "{template['opening']}"
2. Discusses the provided content naturally with back-and-forth conversation
3. Includes natural speech patterns (pauses, "um", "you know", etc.)
4. Has the host asking insightful questions
5. Has the guest providing detailed, engaging answers
6. Includes natural transitions and reactions
7. Ends with: "{template['closing']}"

Format your response as:
HOST: [dialogue]
GUEST: [dialogue]
HOST: [dialogue]
etc.

Keep each segment to 1-2 sentences for natural pacing."""

        if custom_prompt:
            base_prompt += f"\n\nAdditional instructions: {custom_prompt}"
        
        return base_prompt
    
    def _parse_ai_dialogue(self, dialogue_text: str) -> List[Dict[str, str]]:
        """Parsa LiteLLM dialogue response till segments"""
        segments = []
        lines = dialogue_text.strip().split('\n')
        
        for line in lines:
            line = line.strip()
            if line.startswith('HOST:'):
                text = line[5:].strip()
                if text:
                    segments.append({"speaker": "host", "text": text})
            elif line.startswith('GUEST:'):
                text = line[6:].strip()
                if text:
                    segments.append({"speaker": "guest", "text": text})
        
        return segments if segments else [{"speaker": "host", "text": "Welcome to our podcast discussion!"}]
    
    async def _generate_pause(self, duration: float) -> bytes:
        """Generera kort tystnad/paus"""
        return b""
    
    def _structure_text_to_dialogue(
        self, 
        text: str, 
        speakers: List[str]
    ) -> List[Dict[str, str]]:
        """Strukturera text till dialogue mellan talare"""
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        
        dialogue_segments = []
        current_speaker_idx = 0
        
        if len(speakers) >= 2:
            intro_text = f"Välkommen till dagens podcast! Idag ska vi prata om ett mycket intressant ämne."
            dialogue_segments.append({
                "speaker": speakers[0],
                "text": intro_text
            })
        
        for i, paragraph in enumerate(paragraphs):
            speaker = speakers[current_speaker_idx % len(speakers)]
            
            if current_speaker_idx % 2 == 0:
                if i == 0:
                    dialogue_text = f"Låt mig börja med att förklara: {paragraph}"
                elif i == len(paragraphs) - 1:
                    dialogue_text = f"Avslutningsvis vill jag säga att {paragraph}"
                else:
                    dialogue_text = f"Det är viktigt att förstå att {paragraph}. Vad tänker du om det här?"
            else:
                dialogue_text = f"Det är en mycket bra poäng. {paragraph}. Jag skulle också vilja tillägga att detta verkligen visar på komplexiteten i ämnet."
            
            dialogue_segments.append({
                "speaker": speaker,
                "text": dialogue_text
            })
            
            current_speaker_idx += 1
        
        if len(speakers) >= 2:
            dialogue_segments.append({
                "speaker": speakers[0],
                "text": "Tack för ett mycket intressant samtal! Det var verkligen lärorikt att diskutera det här ämnet med dig."
            })
            dialogue_segments.append({
                "speaker": speakers[1],
                "text": "Tack så mycket! Det var ett nöje att vara med och jag hoppas att lyssnarna fann det lika intressant som vi gjorde."
            })
        
        return dialogue_segments
    
    def _combine_audio_segments(self, audio_segments: List[bytes]) -> bytes:
        """Kombinera flera audio segments till en enda fil"""
        combined = b""
        
        for segment in audio_segments:
            combined += segment
        
        return combined
    
    def get_conversational_voices(self) -> Dict[str, Any]:
        """Hämta tillgängliga conversational AI röster"""
        voices_info = self.list_available_voices()
        
        return {
            "available_voices": self.VOICE_IDS,
            "conversation_pairs": self.CONVERSATION_PAIRS,
            "all_voices": voices_info.get("voices", {}),
            "supported_styles": list(self.CONVERSATION_PAIRS.keys())
        }
    
    async def convert_text_to_speech_with_voice_id(
        self,
        text: str,
        voice_id: str,
        language: str = "sv",
        model: str = "eleven_multilingual_v2"
    ) -> bytes:
        """Konvertera text till tal med specifik voice_id"""
        if not text or not text.strip():
            raise ValueError("Text kan inte vara tom")
        if language not in self.SUPPORTED_LANGUAGES:
            raise ValueError(f"Språk '{language}' stöds ej. Använd: {', '.join(self.SUPPORTED_LANGUAGES)}")
        
        api_key = self.api_key
        headers = {
            "xi-api-key": api_key,
            "Content-Type": "application/json"
        }
        url = f"https://api.elevenlabs.io/v1/text-to-speech/{voice_id}"
        data = {
            "text": text,
            "model_id": model,
            "voice_settings": {},
            "output_format": "mp3_44100_128"
        }
        
        async with httpx.AsyncClient(timeout=60) as client:
            response = await client.post(url, headers=headers, json=data)
            if response.status_code != 200:
                raise Exception(f"TTS API error: {response.status_code} {response.text}")
            return response.content
    
    def generate_podcast_genfm(
        self,
        source_type: str,
        source_value: str,
        format: str = "conversation",
        host_voice_id: str = "",
        guest_voice_id: str = "",
        language: str = "sv",
        prompt: str = None,
        model: str = "eleven_v3_alpha"
    ) -> bytes:
        """Generate a podcast using ElevenLabs GenFM/conversational API (synkron wrapper)"""
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        try:
            return loop.run_until_complete(
                self.generate_podcast_genfm_async(
                    source_type, source_value, format, host_voice_id,
                    guest_voice_id, language, prompt, model
                )
            )
        finally:
            loop.close()
    
    async def generate_podcast_genfm_async(
        self,
        source_type: str,
        source_value: str,
        format: str = "conversation",
        host_voice_id: str = "",
        guest_voice_id: str = "",
        language: str = "sv",
        prompt: str = None,
        model: str = "eleven_v3_alpha"
    ) -> bytes:
        """Generate a podcast using ElevenLabs GenFM/conversational API (asynkron)"""
        api_key = self.api_key
        headers = {
            "xi-api-key": api_key,
            "Content-Type": "application/json"
        }
        url = "https://api.elevenlabs.io/v1/studio/podcasts"
        data = {
            "model_id": model,
            "quality_preset": "standard",
            "duration_scale": "default",
            "language": language
        }
        
        if format == "conversation":
            if not host_voice_id or not guest_voice_id:
                raise ValueError("Both host_voice_id and guest_voice_id are required for conversation format.")
            data["mode"] = {
                "type": "conversation",
                "conversation": {
                    "host_voice_id": host_voice_id,
                    "guest_voice_id": guest_voice_id
                }
            }
        else:
            if not host_voice_id:
                raise ValueError("host_voice_id is required for bulletin format.")
            data["mode"] = {
                "type": "bulletin", 
                "bulletin": {
                    "host_voice_id": host_voice_id
                }
            }
        
        if prompt:
            data["highlights"] = prompt
        
        if source_type == "url":
            data["source"] = {"type": "url", "url": source_value}
        elif source_type == "file":
            data["source"] = {"type": "text", "text": source_value}
        elif source_type == "text":
            data["source"] = {"type": "text", "text": source_value}
        else:
            raise ValueError("Invalid source_type. Must be 'url', 'file', or 'text'.")
        
        async with httpx.AsyncClient(timeout=120) as client:
            response = await client.post(url, headers=headers, json=data)
            if response.status_code != 200:
                raise Exception(f"GenFM API error: {response.status_code} {response.text}")
            return response.content


# === INITIALIZE MODULES ===

tts = TextToSpeech()
file_handler = FileHandler()
url_handler = URLHandler()


# === FLASK ROUTES ===

@app.route('/')
def home():
    return send_file('index.html')

@app.route('/health')
def health():
    return jsonify({
        "status": "ok",
        "modules": {
            "tts": "active",
            "file_handler": "active",
            "url_handler": "active",
            "podcast_ai": "active",
            "litellm": "active" if LITELLM_API_KEY else "unavailable"
        },
        "supported_formats": list(file_handler.supported_formats.keys()),
        "features": {
            "text_to_speech": "available",
            "file_to_speech": "available", 
            "url_to_speech": "available",
            "podcast_generation": "available",
            "conversational_ai": "available" if LITELLM_API_KEY else "limited"
        },
        "conversation_styles": list(tts.CONVERSATION_PAIRS.keys())
    })

@app.route('/api/tts', methods=['POST'])
def text_to_speech():
    """Text-till-tal endpoint - Direkt textinput"""
    try:
        data = request.get_json()
        text = data.get('text', '')
        voice = data.get('voice', 'astrid')
        language = data.get('language', 'sv')
        
        if not text:
            return jsonify({"error": "Text krävs"}), 400
        
        print(f"📨 TTS Request: {text[:50]}...")
        
        audio_data = tts.convert_text_to_speech(
            text=text,
            voice=voice,
            language=language
        )
        
        print(f"✅ Genererat {len(audio_data)} bytes")
        
        return Response(
            audio_data,
            mimetype='audio/mpeg',
            headers={
                'Content-Disposition': f'attachment; filename=tal_{language}.mp3'
            }
        )
        
    except Exception as e:
        print(f"❌ FEL: {type(e).__name__}: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/tts/file', methods=['POST'])
def file_to_speech():
    """Konvertera fil (PDF, DOCX, TXT, etc.) till tal"""
    try:
        if 'file' not in request.files:
            return jsonify({"error": "Ingen fil uppladdad"}), 400
        
        file = request.files['file']
        voice = request.form.get('voice', 'astrid')
        language = request.form.get('language', 'sv')
        
        if file.filename == '':
            return jsonify({"error": "Ingen fil vald"}), 400
        
        file_ext = Path(file.filename).suffix.lower()
        
        if file_ext not in file_handler.supported_formats:
            supported = ", ".join(file_handler.supported_formats.keys())
            return jsonify({
                "error": f"Filformat '{file_ext}' stöds ej. Stödda format: {supported}"
            }), 400
        
        # Spara fil temporärt
        temp_path = UPLOAD_DIR / f"temp_{file.filename}"
        file.save(str(temp_path))
        
        try:
            print(f"📄 Sparad fil: {temp_path}")
            
            # Extrahera text
            file_result = file_handler.extract_text(str(temp_path))
            text = file_result["text"]
            
            print(f"📝 Extraherade {file_result['word_count']} ord från {file_result['file_type']}")
            
            # Begränsa textlängd
            if len(text) > 5000:
                text = text[:5000] + "... (förkortat för demo)"
                print("⚠️ Text förkortad till 5000 tecken")
            
            # Konvertera till tal
            audio_data = tts.convert_text_to_speech(
                text=text,
                voice=voice,
                language=language
            )
            
            print(f"✅ {file_result['file_type']} → TTS klar! {len(audio_data)} bytes")
            
            return Response(
                audio_data,
                mimetype='audio/mpeg',
                headers={
                    'Content-Disposition': f'attachment; filename={Path(file.filename).stem}_tal.mp3'
                }
            )
            
        finally:
            # Rensa temp-fil
            if temp_path.exists():
                temp_path.unlink()
                
    except Exception as e:
        print(f"❌ Fel vid filkonvertering: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/file/extract', methods=['POST'])
def extract_file_text():
    """Extrahera text från valfri fil"""
    try:
        if 'file' not in request.files:
            return jsonify({"error": "Ingen fil uppladdad"}), 400
        
        file = request.files['file']
        
        if file.filename == '':
            return jsonify({"error": "Ingen fil vald"}), 400
        
        file_ext = Path(file.filename).suffix.lower()
        
        if file_ext not in file_handler.supported_formats:
            supported = ", ".join(file_handler.supported_formats.keys())
            return jsonify({
                "error": f"Filformat '{file_ext}' stöds ej. Stödda format: {supported}"
            }), 400
        
        # Spara fil temporärt
        temp_path = UPLOAD_DIR / f"temp_{file.filename}"
        file.save(str(temp_path))
        
        try:
            # Extrahera text
            file_result = file_handler.extract_text(str(temp_path))
            
            return jsonify({
                "success": True,
                "text": file_result["text"],
                "paragraphs": file_result["paragraphs"],
                "statistics": {
                    "word_count": file_result["word_count"],
                    "character_count": file_result["character_count"],
                    "file_type": file_result["file_type"],
                    "file_size_mb": file_result["file_size_mb"]
                },
                "metadata": file_result.get("metadata", {}),
                "tts_estimate": tts.estimate_cost(file_result["text"])
            })
            
        finally:
            # Rensa temp-fil
            if temp_path.exists():
                temp_path.unlink()
                
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/tts/url', methods=['POST'])
def url_to_speech():
    """Konvertera URL (webbsida eller PDF) till tal"""
    try:
        data = request.get_json()
        url = data.get('url', '')
        voice = data.get('voice', 'astrid')
        language = data.get('language', 'sv')
        
        if not url:
            return jsonify({"error": "URL krävs"}), 400
        
        print(f"🌐 URL TTS Request: {url}")
        
        # Validera URL
        validation = url_handler.validate_url(url)
        if not validation["valid"]:
            return jsonify({"error": f"Ogiltig URL: {validation['error']}"}), 400
        
        if not validation.get("accessible", True):
            return jsonify({"error": "URL:en är inte tillgänglig"}), 400
        
        # Extrahera text
        url_result = url_handler.extract_from_url(url)
        text = url_result["text"]
        
        print(f"📝 Extraherade {url_result['word_count']} ord från {url_result['source_type']}")
        
        # Begränsa textlängd
        if len(text) > 5000:
            text = text[:5000] + "... (förkortat för demo)"
            print("⚠️ Text förkortad till 5000 tecken")
        
        # Konvertera till tal
        audio_data = tts.convert_text_to_speech(
            text=text,
            voice=voice,
            language=language
        )
        
        print(f"✅ URL → TTS klar! {len(audio_data)} bytes")
        
        # Skapa filnamn
        url_title = url_result.get("metadata", {}).get("title", "url_content")
        safe_title = re.sub(r'[^\w\s-]', '', url_title)[:50]
        
        return Response(
            audio_data,
            mimetype='audio/mpeg',
            headers={
                'Content-Disposition': f'attachment; filename={safe_title}_tal.mp3'
            }
        )
        
    except Exception as e:
        print(f"❌ Fel vid URL-konvertering: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/url/extract', methods=['POST'])
def extract_url_text():
    """Extrahera text från URL"""
    try:
        data = request.get_json()
        url = data.get('url', '')
        
        if not url:
            return jsonify({"error": "URL saknas"}), 400
        
        # Validera URL
        validation = url_handler.validate_url(url)
        if not validation["valid"]:
            return jsonify({"error": f"Ogiltig URL: {validation['error']}"}), 400
        
        if not validation.get("accessible", True):
            return jsonify({"error": "URL:en är inte tillgänglig"}), 400
        
        # Extrahera text
        url_result = url_handler.extract_from_url(url)
        
        return jsonify({
            "success": True,
            "text": url_result["text"],
            "paragraphs": url_result["paragraphs"],
            "statistics": {
                "word_count": url_result["word_count"],
                "character_count": url_result["character_count"],
                "source_type": url_result["source_type"],
                "source_url": url_result["source_url"]
            },
            "metadata": url_result.get("metadata", {}),
            "tts_estimate": tts.estimate_cost(url_result["text"])
        })
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/podcast/enhanced', methods=['POST'])
def generate_enhanced_podcast():
    """Generate an enhanced podcast using LiteLLM + ElevenLabs"""
    try:
        data = request.get_json()
        
        source_type = data.get('source_type')
        source_value = data.get('source_value')
        template_style = data.get('template_style', 'casual_chat')
        host_voice_id = data.get('host_voice_id')
        guest_voice_id = data.get('guest_voice_id')
        language = data.get('language', 'en')
        custom_prompt = data.get('custom_prompt', '')
        model = data.get('model', 'eleven_multilingual_v2')
        
        print(f"🎙️ Enhanced Podcast Request: {template_style}")
        
        # Hämta källtext
        source_text = ""
        if source_type == "text":
            source_text = source_value
        elif source_type == "url":
            url_result = url_handler.extract_from_url(source_value)
            source_text = url_result["text"]
        elif source_type == "file":
            source_text = source_value
        else:
            return jsonify({"error": "Invalid source_type"}), 400
        
        if not source_text.strip():
            return jsonify({"error": "No content found to generate podcast from"}), 400
        
        # Begränsa textlängd
        if len(source_text) > 4000:
            source_text = source_text[:4000] + "..."
            print("⚠️ Text förkortad till 4000 tecken för podcast")
        
        # Generera enhanced conversation
        audio_data = tts.generate_enhanced_conversation(
            text=source_text,
            template_style=template_style,
            host_voice_id=host_voice_id,
            guest_voice_id=guest_voice_id,
            language=language,
            custom_prompt=custom_prompt,
            model=model
        )
        
        filename = f"enhanced_podcast_{template_style}_{language}.mp3"
        
        return Response(
            audio_data,
            mimetype='audio/mpeg',
            headers={
                'Content-Disposition': f'attachment; filename={filename}'
            }
        )
        
    except Exception as e:
        print(f"❌ Enhanced podcast error: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/podcast/templates')
def get_conversation_templates():
    """Get available conversation templates"""
    return jsonify({
        "templates": tts._conversation_templates,
        "models": tts.AVAILABLE_MODELS
    })

@app.route('/api/voices')
def list_voices():
    """Lista tillgängliga röster"""
    return jsonify(tts.list_available_voices())

@app.route('/api/podcast/voices')
def list_podcast_voices():
    """Lista tillgängliga röster för podcast-generering"""
    return jsonify(tts.get_conversational_voices())

@app.route('/api/supported-formats')
def get_supported_formats():
    """Lista alla filformat som stöds"""
    return jsonify({
        "formats": list(file_handler.supported_formats.keys()),
        "details": {
            ".pdf": "Adobe PDF dokument",
            ".docx": "Microsoft Word dokument",
            ".doc": "Microsoft Word (legacy) - begränsat stöd",
            ".txt": "Ren textfil",
            ".md": "Markdown dokument",
            ".rtf": "Rich Text Format - begränsat stöd"
        }
    })


# === CLEANUP AND STARTUP ===

def cleanup_old_uploads():
    """Ta bort gamla temp-filer"""
    for file in UPLOAD_DIR.glob("temp*"):
        try:
            file.unlink()
        except:
            pass

cleanup_old_uploads()


# === MAIN ===

if __name__ == "__main__":
    print("🚀 Startar Formatmaskinen API med Flask + LiteLLM...")
    print(f"📁 Arbetskatalog: {os.getcwd()}")
    print(f"🔑 ElevenLabs API-nyckel finns: {'Ja' if os.getenv('ELEVENLABS_API_KEY') else 'NEJ!'}")
    print(f"🤖 LiteLLM API-nyckel finns: {'Ja' if LITELLM_API_KEY else 'NEJ!'}")
    print(f"🌐 LiteLLM Base URL: {LITELLM_BASE_URL}")
    print(f"📄 Filstöd: {', '.join(file_handler.supported_formats.keys())}")
    
    app.run(host='0.0.0.0', port=8000, debug=True)